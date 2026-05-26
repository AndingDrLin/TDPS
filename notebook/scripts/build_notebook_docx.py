from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from notebook_content import ENTRIES

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_CM = 1.8


def _set_layout(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)


def _apply_default_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = 1.15


def _add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _set_footer_page_numbers(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        r = p.add_run("Page ")
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)
        _add_field(p, "PAGE")
        r = p.add_run(" of ")
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)
        _add_field(p, "NUMPAGES")


def _add_line(document: Document, text: str, bold: bool = False, size: float = 10.0) -> None:
    p = document.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def _add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(item)


def _add_code_block(document: Document, lines: list[str], language: str) -> None:
    _add_line(document, f"Core snippet ({language}):", bold=True)
    for line in lines:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9)


def _add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)

    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            run = cell.paragraphs[0].runs[0]
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)


def _add_image(document: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        raise FileNotFoundError(f"Missing figure: {image_path}")

    document.add_picture(str(image_path), width=Inches(6.3))
    p = document.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = document.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)


def _add_entry_page(document: Document, entry: dict, image_dir: Path) -> None:
    _add_line(document, f"Date: {entry['date']} | Page: {entry['page']}", bold=True, size=11)
    _add_line(document, f"Title: {entry['title']}", bold=True, size=11)
    _add_line(document, f"Purpose: {entry['purpose']}")

    _add_line(document, "1. Brief Procedure:", bold=True)
    _add_bullets(document, entry["brief_procedure"])

    _add_line(document, "2. Workings & Calculations:", bold=True)
    _add_bullets(document, entry["workings"])

    _add_line(document, "3. Engineering Struggle & Code Implementation:", bold=True)
    _add_bullets(document, entry["engineering_struggle"])
    _add_code_block(document, entry["code_snippet"], entry["code_snippet_language"])

    _add_line(document, "4. Observations and Recorded Data:", bold=True)
    _add_table(document, entry["observations_table"]["headers"], entry["observations_table"]["rows"])

    fig = entry["figure"]
    _add_image(document, image_dir / fig["file"], fig["title"])

    _add_line(document, "5. Conclusion & Next Steps:", bold=True)
    _add_line(document, entry["conclusion"])


def build_notebook_docx(base_dir: Path, output_path: Path | None = None) -> Path:
    image_dir = base_dir / "assets" / "images"
    output_path = output_path or (base_dir / "outputs" / "team15_ptsd_lab_notebook.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_layout(doc)
    _apply_default_style(doc)

    for idx, entry in enumerate(ENTRIES):
        _add_entry_page(doc, entry, image_dir)
        if idx != len(ENTRIES) - 1:
            doc.add_page_break()

    _set_footer_page_numbers(doc)
    doc.save(output_path)
    print(f"Generated DOCX: {output_path}")
    return output_path


if __name__ == "__main__":
    root = Path(__file__).resolve().parent.parent
    build_notebook_docx(root)
