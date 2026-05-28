"""Build TDPS Initial Design Report in DOCX format with fixed 5-page structure."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from prepare_schematic_crops import prepare_schematic_crops
from report_content import (
    COURSE_NAME,
    CURRENT_DATE,
    HARDWARE_TABLE_HEADERS,
    HARDWARE_TABLE_ROWS,
    INITIAL_TEST_TABLE_HEADERS,
    INITIAL_TEST_TABLE_ROWS,
    MAX_TECHNICAL_WORDS,
    PAGE2_CALCULATION,
    PAGE2_HEADING,
    PAGE2_PARAGRAPHS,
    PAGE3_HEADING,
    PAGE3_PARAGRAPHS,
    PAGE4_HEADING,
    PAGE4_PARAGRAPHS,
    PAGE5_HEADING,
    PAGE5_PARAGRAPH,
    REPORT_TITLE,
    TEAM_MEMBERS,
    TEAM_NAME,
    TEAM_NUMBER,
    technical_word_count,
)

NAVY_RGB = RGBColor(0x00, 0x33, 0x66)

# Page layout: A4 with 1.4 cm margins  ->  18.2 cm usable width
PAGE_WIDTH_CM = 21.0
MARGIN_LR_CM = 1.4
MARGIN_TB_CM = 1.4


# ── low-level helpers ───────────────────────────────────────────────────────

def _set_document_layout(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(MARGIN_LR_CM)
    section.right_margin = Cm(MARGIN_LR_CM)
    section.top_margin = Cm(MARGIN_TB_CM)
    section.bottom_margin = Cm(MARGIN_TB_CM)


def _apply_default_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1.5)
    normal.paragraph_format.line_spacing = 1.0


def _set_cell_fill(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_cell_margins(cell, top: int = 30, start: int = 60, bottom: int = 30, end: int = 60) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_text(cell, text: str, bold: bool = False, white: bool = False, size: float = 8.5) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _set_footer_page_numbers(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.clear()
        r = paragraph.add_run("Page ")
        r.font.size = Pt(8)
        r.font.name = "Times New Roman"
        _add_field(paragraph, "PAGE")
        r = paragraph.add_run(" of ")
        r.font.size = Pt(8)
        r.font.name = "Times New Roman"
        _add_field(paragraph, "NUMPAGES")


def _remove_table_borders(table) -> None:
    """Remove all borders from a table (used for image layout tables)."""
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl_pr.append(borders)


# ── document building blocks ────────────────────────────────────────────────

def _add_heading(document: Document, text: str, size: int = 12) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Helvetica"
    run.font.size = Pt(size)
    run.font.color.rgb = NAVY_RGB


def _add_body_text(document: Document, text: str, size: float = 9.5) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def _add_caption(document: Document, text: str) -> None:
    cap = document.add_paragraph(text)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(1)
    cap.paragraph_format.space_after = Pt(2)
    run = cap.runs[0]
    run.font.size = Pt(8)
    run.italic = True


def _add_horizontal_rule(document: Document, color: str = "003366") -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _insert_image(document: Document, image_path: Path, caption: str, width_inches: float = 6.5) -> None:
    if not image_path.exists():
        raise FileNotFoundError(f"Missing image: {image_path}")
    document.add_picture(str(image_path), width=Inches(width_inches))
    last = document.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    last.paragraph_format.space_before = Pt(2)
    last.paragraph_format.space_after = Pt(0)
    _add_caption(document, caption)


def _insert_side_by_side_images(
    document: Document,
    left_image: Path,
    right_image: Path,
    left_caption: str,
    right_caption: str,
    width_inches: float = 3.4,
) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _remove_table_borders(table)

    for cell, img, caption in [
        (table.cell(0, 0), left_image, left_caption),
        (table.cell(0, 1), right_image, right_caption),
    ]:
        if not img.exists():
            raise FileNotFoundError(f"Missing image: {img}")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p_img = cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(str(img), width=Inches(width_inches))

        p_cap = cell.add_paragraph(caption)
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(1)
        p_cap.paragraph_format.space_after = Pt(0)
        cap_run = p_cap.runs[0]
        cap_run.italic = True
        cap_run.font.size = Pt(7.5)


def _add_styled_table(
    document: Document,
    headers: list[str],
    rows: list[tuple[str, ...]],
    col_widths_cm: list[float],
) -> None:
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    for idx, width_cm in enumerate(col_widths_cm):
        table.columns[idx].width = Cm(width_cm)
        for row in table.rows:
            cell = row.cells[idx]
            cell.width = Cm(width_cm)
            _set_cell_margins(cell)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _set_cell_fill(cell, "003366")
        _set_cell_text(cell, header, bold=True, white=True, size=8.2)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for r_idx, row_data in enumerate(rows, start=1):
        row = table.rows[r_idx]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.68)
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            row_fill = "F2F2F2" if r_idx % 2 == 1 else "FFFFFF"
            _set_cell_fill(cell, row_fill)
            _set_cell_text(cell, value, size=8.0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


# ── page builders ───────────────────────────────────────────────────────────

def _add_cover_page(document: Document, crops: dict[str, Path]) -> None:
    for _ in range(3):
        document.add_paragraph("")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(REPORT_TITLE)
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.name = "Helvetica"
    tr.font.color.rgb = NAVY_RGB

    _add_horizontal_rule(document)

    course = document.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = course.add_run(f"Course: {COURSE_NAME}")
    r.font.size = Pt(13)
    r.font.name = "Times New Roman"

    team = document.add_paragraph()
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = team.add_run(f"Team: {TEAM_NAME}  |  Team Number: {TEAM_NUMBER}")
    r.font.size = Pt(11)

    document.add_paragraph("")
    mh = document.add_paragraph()
    mh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = mh.add_run("Team Members")
    r.bold = True
    r.font.size = Pt(11)

    for name, guid in TEAM_MEMBERS:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{name}  (GUID: {guid})")
        r.font.size = Pt(10.5)

    document.add_paragraph("")
    dp = document.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dp.add_run(f"Date: {CURRENT_DATE}")
    r.font.size = Pt(11)

    # PCB V2 layout preview fills the otherwise-empty lower half of the cover
    pcb = crops.get("pcb2_layout")
    if pcb and pcb.exists():
        document.add_paragraph("")
        _insert_image(document, pcb, "PCB V2 Layout Preview", width_inches=3.2)


def _add_page2(document: Document, image_dir: Path, crops: dict[str, Path]) -> None:
    _add_heading(document, PAGE2_HEADING)
    for paragraph in PAGE2_PARAGRAPHS:
        _add_body_text(document, paragraph)

    _insert_image(
        document,
        image_dir / "block_diagram.png",
        "Figure 1. System architecture with separated power and data/control paths.",
        width_inches=6.2,
    )

    # V1 vs V2 power schematic comparison — directly supports architecture discussion
    _insert_side_by_side_images(
        document,
        crops["v1_power"],
        crops["v2_power"],
        "Figure 1a. Schematic V1: power input baseline",
        "Figure 1b. Schematic V2: enhanced protection",
        width_inches=3.4,
    )

    # Hardware table: col widths sum to ~18.2 cm (full usable width)
    _add_styled_table(document, HARDWARE_TABLE_HEADERS, HARDWARE_TABLE_ROWS, [2.6, 4.6, 11.0])

    calc = document.add_paragraph()
    calc.paragraph_format.space_before = Pt(2)
    calc_run = calc.add_run("Calculation: " + PAGE2_CALCULATION)
    calc_run.bold = True
    calc_run.font.size = Pt(9.0)


def _add_page3(document: Document, image_dir: Path, crops: dict[str, Path]) -> None:
    _add_heading(document, PAGE3_HEADING)
    for paragraph in PAGE3_PARAGRAPHS:
        _add_body_text(document, paragraph)

    _insert_image(
        document,
        image_dir / "flowchart.png",
        "Figure 2. Software control flow for PID, obstacle handling, and telemetry.",
        width_inches=3.2,
    )

    # V2 motor driver schematic evidence — shows DRV8874 implementation detail
    v2_motor = crops.get("v2_motor")
    if v2_motor and v2_motor.exists():
        _insert_image(
            document,
            v2_motor,
            "Figure 3. Schematic V2: Dual DRV8874PWPR motor driver with encoder feedback.",
            width_inches=5.5,
        )


def _add_page4(document: Document, crops: dict[str, Path]) -> None:
    _add_heading(document, PAGE4_HEADING)
    for paragraph in PAGE4_PARAGRAPHS:
        _add_body_text(document, paragraph)

    # V1 vs V2 motor driver comparison — TB6612 (V1) vs DRV8874 (V2)
    _insert_side_by_side_images(
        document,
        crops["v1_motor"],
        crops["v2_motor"],
        "Figure 4a. V1: TB6612 motor driver",
        "Figure 4b. V2: Dual DRV8874PWPR drivers",
        width_inches=3.4,
    )

    # PCB V2 layout — physical evidence of routed board
    pcb = crops.get("pcb2_layout")
    if pcb and pcb.exists():
        _insert_image(
            document,
            pcb,
            "Figure 5. PCB V2 routed layout with separated power and signal domains.",
            width_inches=3.8,
        )

    # Test data table: col widths sum to ~18.2 cm
    _add_styled_table(
        document, INITIAL_TEST_TABLE_HEADERS, INITIAL_TEST_TABLE_ROWS, [3.6, 2.4, 12.2]
    )


def _add_page5(document: Document, image_dir: Path, crops: dict[str, Path]) -> None:
    _add_heading(document, PAGE5_HEADING)

    _insert_image(
        document,
        image_dir / "gantt_chart.png",
        "Figure 6. Week 1-15 Gantt chart with current-week marker and milestone timeline.",
        width_inches=6.9,
    )
    _add_body_text(document, PAGE5_PARAGRAPH)

    # V2 interface connector schematic — fills remaining page space with evidence
    v2_intf = crops.get("v2_interface")
    if v2_intf and v2_intf.exists():
        _insert_image(
            document,
            v2_intf,
            "Figure 7. Schematic V2: Interface connectors for LoRa, Radar, and OLED modules.",
            width_inches=4.0,
        )


# ── main entry ──────────────────────────────────────────────────────────────

def build_docx_report(base_dir: Path, image_dir: Path, output_path: Path | None = None) -> Path:
    """Build and save the DOCX report with fixed page-by-page content and schematic evidence."""
    word_count = technical_word_count()
    if word_count > MAX_TECHNICAL_WORDS:
        raise ValueError(f"Technical content exceeds {MAX_TECHNICAL_WORDS} words: {word_count}")

    output_path = output_path or (base_dir / "outputs" / "final" / "tdps_initial_design_report_generated.docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    crops = prepare_schematic_crops(image_dir)

    document = Document()
    _set_document_layout(document)
    _apply_default_styles(document)

    _add_cover_page(document, crops)
    document.add_page_break()

    _add_page2(document, image_dir, crops)
    document.add_page_break()

    _add_page3(document, image_dir, crops)
    document.add_page_break()

    _add_page4(document, crops)
    document.add_page_break()

    _add_page5(document, image_dir, crops)

    _set_footer_page_numbers(document)
    document.save(output_path)

    print(f"Technical word count: {word_count}")
    print(f"DOCX generated: {output_path}")
    return output_path


if __name__ == "__main__":
    root = Path(__file__).resolve().parent.parent
    build_docx_report(root, root / "assets" / "images")
